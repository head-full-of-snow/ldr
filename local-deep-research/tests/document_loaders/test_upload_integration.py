"""
Integration tests for document loaders in the upload workflow.

These tests create real temporary files and verify the extraction works
end-to-end for various file formats.
"""

import pytest


class TestRealFileExtraction:
    """Tests that create real files and verify extraction works."""

    def test_txt_file_extraction(self, tmp_path):
        """Test extracting text from a real .txt file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real txt file
        content = (
            "This is a test text file.\nWith multiple lines.\nAnd some content."
        )
        file_path = tmp_path / "test.txt"
        file_path.write_text(content)

        # Read as bytes (simulating upload)
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".txt", "test.txt")

        assert extracted is not None
        assert "test text file" in extracted
        assert "multiple lines" in extracted

    def test_json_file_extraction(self, tmp_path):
        """Test extracting text from a real .json file."""
        from local_deep_research.document_loaders import extract_text_from_bytes
        import json

        # Create a real json file
        data = {
            "title": "Research Document",
            "author": "Test Author",
            "sections": [
                {
                    "name": "Introduction",
                    "content": "This is the introduction.",
                },
                {"name": "Methods", "content": "These are the methods."},
            ],
        }
        file_path = tmp_path / "document.json"
        file_path.write_text(json.dumps(data, indent=2))

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(
            file_bytes, ".json", "document.json"
        )

        assert extracted is not None
        assert "Research Document" in extracted
        assert "Test Author" in extracted
        assert "Introduction" in extracted
        assert "Methods" in extracted

    def test_yaml_file_extraction(self, tmp_path):
        """Test extracting text from a real .yaml file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real yaml file
        yaml_content = """
title: Configuration File
database:
  host: localhost
  port: 5432
  name: research_db
settings:
  debug: true
  log_level: INFO
"""
        file_path = tmp_path / "config.yaml"
        file_path.write_text(yaml_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".yaml", "config.yaml")

        assert extracted is not None
        assert "Configuration File" in extracted
        assert "localhost" in extracted
        assert "research_db" in extracted

    def test_yml_file_extraction(self, tmp_path):
        """Test extracting text from a real .yml file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real yml file
        yml_content = """
project: Local Deep Research
version: 1.0.0
features:
  - document_loaders
  - rag_indexing
  - web_search
"""
        file_path = tmp_path / "project.yml"
        file_path.write_text(yml_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".yml", "project.yml")

        assert extracted is not None
        assert "Local Deep Research" in extracted
        assert "document_loaders" in extracted

    def test_md_file_extraction(self, tmp_path):
        """Test extracting text from a real .md file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real markdown file
        md_content = """# Research Notes

## Introduction

This document contains research notes about document loaders.

## Key Findings

1. JSON and YAML are well supported
2. PDF extraction works with PyPDF
3. HTML parsing handles nested elements

## Conclusion

The document loaders module provides comprehensive format support.
"""
        file_path = tmp_path / "notes.md"
        file_path.write_text(md_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".md", "notes.md")

        assert extracted is not None
        assert "Research Notes" in extracted or "document loaders" in extracted

    def test_html_file_extraction(self, tmp_path):
        """Test extracting text from a real .html file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real HTML file
        html_content = """<!DOCTYPE html>
<html>
<head>
    <title>Test Document</title>
</head>
<body>
    <h1>Welcome to the Test Page</h1>
    <p>This is a paragraph with some <strong>important</strong> content.</p>
    <ul>
        <li>First item</li>
        <li>Second item</li>
    </ul>
    <div class="content">
        <p>More detailed information here.</p>
    </div>
</body>
</html>
"""
        file_path = tmp_path / "page.html"
        file_path.write_text(html_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".html", "page.html")

        assert extracted is not None
        # HTML text extraction should get the content
        assert (
            "Welcome" in extracted
            or "Test" in extracted
            or "paragraph" in extracted
        )

    def test_csv_file_extraction(self, tmp_path):
        """Test extracting text from a real .csv file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real CSV file
        csv_content = """name,age,city,occupation
John Doe,35,New York,Engineer
Jane Smith,28,Los Angeles,Designer
Bob Johnson,42,Chicago,Manager
Alice Brown,31,Boston,Developer
"""
        file_path = tmp_path / "data.csv"
        file_path.write_text(csv_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".csv", "data.csv")

        assert extracted is not None
        assert "John Doe" in extracted or "name" in extracted

    def test_tsv_file_extraction(self, tmp_path):
        """Test extracting text from a real .tsv file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real TSV file
        tsv_content = "name\tvalue\tdescription\nitem1\t100\tFirst item\nitem2\t200\tSecond item"
        file_path = tmp_path / "data.tsv"
        file_path.write_text(tsv_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".tsv", "data.tsv")

        assert extracted is not None
        assert "item1" in extracted or "name" in extracted

    def test_xml_file_extraction(self, tmp_path):
        """Test extracting text from a real .xml file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real XML file
        xml_content = """<?xml version="1.0" encoding="UTF-8"?>
<document>
    <title>Research Paper</title>
    <author>Jane Researcher</author>
    <abstract>
        This paper discusses the implementation of document loaders
        for various file formats in the Local Deep Research project.
    </abstract>
    <sections>
        <section id="1">
            <heading>Introduction</heading>
            <content>The introduction section content goes here.</content>
        </section>
        <section id="2">
            <heading>Methodology</heading>
            <content>The methodology section describes our approach.</content>
        </section>
    </sections>
</document>
"""
        file_path = tmp_path / "document.xml"
        file_path.write_text(xml_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text
        extracted = extract_text_from_bytes(file_bytes, ".xml", "document.xml")

        assert extracted is not None
        assert "Research Paper" in extracted or "Jane Researcher" in extracted

    def test_toml_file_extraction(self, tmp_path):
        """Test extracting text from a real .toml file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real TOML file
        toml_content = """
[project]
name = "local-deep-research"
version = "1.3.0"
description = "AI-powered research assistant"

[dependencies]
flask = "^2.0"
langchain = "^0.1"

[settings]
debug = true
log_level = "INFO"
"""
        file_path = tmp_path / "config.toml"
        file_path.write_text(toml_content)

        # Read as bytes
        file_bytes = file_path.read_bytes()

        # Extract text - may return None if tomli is not installed
        extracted = extract_text_from_bytes(file_bytes, ".toml", "config.toml")

        # TOML extraction requires optional 'tomli' dependency
        # Skip assertion if it returns None (dependency not installed)
        if extracted is not None:
            assert "local-deep-research" in extracted or "project" in extracted
        else:
            pytest.skip("TOML extraction requires 'tomli' package")


class TestUnicodeFileExtraction:
    """Tests for files with unicode content."""

    def test_txt_with_unicode(self, tmp_path):
        """Test text file with unicode characters."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        content = "Hello 世界! Привет мир! مرحبا بالعالم 🌍🌎🌏"
        file_path = tmp_path / "unicode.txt"
        file_path.write_text(content, encoding="utf-8")

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".txt", "unicode.txt")

        assert extracted is not None
        assert "Hello" in extracted
        assert "世界" in extracted
        assert "Привет" in extracted

    def test_json_with_unicode(self, tmp_path):
        """Test JSON file with unicode content."""
        from local_deep_research.document_loaders import extract_text_from_bytes
        import json

        data = {
            "greeting_en": "Hello World",
            "greeting_zh": "你好世界",
            "greeting_ja": "こんにちは世界",
            "emoji": "Research 📚 Knowledge 🧠",
        }
        file_path = tmp_path / "unicode.json"
        file_path.write_text(
            json.dumps(data, ensure_ascii=False), encoding="utf-8"
        )

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".json", "unicode.json")

        assert extracted is not None
        assert "Hello World" in extracted
        assert "你好世界" in extracted

    def test_yaml_with_unicode(self, tmp_path):
        """Test YAML file with unicode content."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        yaml_content = """
title: 研究ノート
author: 田中太郎
content: This document contains 日本語 and English text.
tags:
  - 研究
  - AI
  - документы
"""
        file_path = tmp_path / "unicode.yaml"
        file_path.write_text(yaml_content, encoding="utf-8")

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".yaml", "unicode.yaml")

        assert extracted is not None
        assert "研究ノート" in extracted
        assert "田中太郎" in extracted


class TestLargeFileExtraction:
    """Tests for handling larger files."""

    def test_large_txt_file(self, tmp_path):
        """Test extraction from a larger text file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a ~100KB text file
        lines = [
            f"Line {i}: This is test content for line number {i}."
            for i in range(2000)
        ]
        content = "\n".join(lines)
        file_path = tmp_path / "large.txt"
        file_path.write_text(content)

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".txt", "large.txt")

        assert extracted is not None
        assert "Line 0:" in extracted
        assert "Line 1999:" in extracted
        assert len(extracted) > 50000  # Should have substantial content

    def test_large_json_file(self, tmp_path):
        """Test extraction from a larger JSON file."""
        from local_deep_research.document_loaders import extract_text_from_bytes
        import json

        # Create a JSON file with many entries
        data = {
            "items": [
                {
                    "id": i,
                    "name": f"Item {i}",
                    "description": f"Description for item {i}",
                }
                for i in range(500)
            ]
        }
        file_path = tmp_path / "large.json"
        file_path.write_text(json.dumps(data, indent=2))

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".json", "large.json")

        assert extracted is not None
        assert "Item 0" in extracted
        assert "Item 499" in extracted


class TestEdgeCases:
    """Tests for edge cases in file extraction."""

    def test_empty_txt_file(self, tmp_path):
        """Test extraction from empty text file."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        file_path = tmp_path / "empty.txt"
        file_path.write_text("")

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".txt", "empty.txt")

        # Should return empty string or None
        assert extracted == "" or extracted is None

    def test_empty_json_file(self, tmp_path):
        """Test extraction from JSON file with empty object."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        file_path = tmp_path / "empty.json"
        file_path.write_text("{}")

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".json", "empty.json")

        # Should handle gracefully
        assert extracted is not None or extracted == ""

    def test_whitespace_only_txt(self, tmp_path):
        """Test extraction from file with only whitespace."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        file_path = tmp_path / "whitespace.txt"
        file_path.write_text("   \n\n\t\t   \n   ")

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(
            file_bytes, ".txt", "whitespace.txt"
        )

        # Should handle gracefully
        assert extracted is not None

    def test_nested_json(self, tmp_path):
        """Test extraction from deeply nested JSON."""
        from local_deep_research.document_loaders import extract_text_from_bytes
        import json

        data = {
            "level1": {
                "level2": {
                    "level3": {
                        "level4": {
                            "level5": {
                                "deep_value": "Found at depth 5",
                                "another": "Deep nested content",
                            }
                        }
                    }
                }
            }
        }
        file_path = tmp_path / "nested.json"
        file_path.write_text(json.dumps(data, indent=2))

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".json", "nested.json")

        assert extracted is not None
        assert "Found at depth 5" in extracted
        assert "Deep nested content" in extracted

    def test_json_array_at_root(self, tmp_path):
        """Test extraction from JSON with array at root."""
        from local_deep_research.document_loaders import extract_text_from_bytes
        import json

        data = [
            {"title": "First Document", "content": "First content"},
            {"title": "Second Document", "content": "Second content"},
            {"title": "Third Document", "content": "Third content"},
        ]
        file_path = tmp_path / "array.json"
        file_path.write_text(json.dumps(data, indent=2))

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".json", "array.json")

        assert extracted is not None
        assert "First Document" in extracted
        assert "Second Document" in extracted
        assert "Third Document" in extracted

    def test_multi_document_yaml(self, tmp_path):
        """Test extraction from multi-document YAML."""
        from local_deep_research.document_loaders import extract_text_from_bytes

        yaml_content = """---
title: First Document
content: Content of first document
---
title: Second Document
content: Content of second document
---
title: Third Document
content: Content of third document
"""
        file_path = tmp_path / "multi.yaml"
        file_path.write_text(yaml_content)

        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".yaml", "multi.yaml")

        assert extracted is not None
        # Should extract from at least some documents
        assert "First Document" in extracted or "Second Document" in extracted


class TestFilePathExtraction:
    """Tests for extracting from files using get_loader_for_path."""

    def test_load_txt_from_path(self, tmp_path):
        """Test loading text file using get_loader_for_path."""
        from local_deep_research.document_loaders import get_loader_for_path

        content = "This is content loaded from a file path."
        file_path = tmp_path / "path_test.txt"
        file_path.write_text(content)

        loader = get_loader_for_path(file_path)

        assert loader is not None
        docs = loader.load()
        assert len(docs) > 0
        assert "loaded from a file path" in docs[0].page_content

    def test_load_json_from_path(self, tmp_path):
        """Test loading JSON file using get_loader_for_path."""
        from local_deep_research.document_loaders import get_loader_for_path
        import json

        data = {"title": "Path Test", "content": "Loaded via path"}
        file_path = tmp_path / "path_test.json"
        file_path.write_text(json.dumps(data))

        loader = get_loader_for_path(file_path)

        assert loader is not None
        docs = loader.load()
        assert len(docs) > 0
        assert "Path Test" in docs[0].page_content

    def test_load_yaml_from_path(self, tmp_path):
        """Test loading YAML file using get_loader_for_path."""
        from local_deep_research.document_loaders import get_loader_for_path

        yaml_content = "title: YAML Path Test\ncontent: Loaded via path API"
        file_path = tmp_path / "path_test.yaml"
        file_path.write_text(yaml_content)

        loader = get_loader_for_path(file_path)

        assert loader is not None
        docs = loader.load()
        assert len(docs) > 0
        assert "YAML Path Test" in docs[0].page_content

    def test_unsupported_path_returns_none(self, tmp_path):
        """Test that unsupported file extension returns None."""
        from local_deep_research.document_loaders import get_loader_for_path

        file_path = tmp_path / "unsupported.xyz"
        file_path.write_text("Some content")

        loader = get_loader_for_path(file_path)

        assert loader is None


class TestBinaryFileFormats:
    """Tests for binary file formats (DOCX, XLSX) that require special libraries."""

    def test_docx_file_extraction(self, tmp_path):
        """Test extracting text from a real .docx file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real DOCX file
        doc = DocxDocument()
        doc.add_heading("Test Document Title", 0)
        doc.add_paragraph("This is a test paragraph in the Word document.")
        doc.add_paragraph("It contains multiple paragraphs for testing.")

        file_path = tmp_path / "test.docx"
        doc.save(str(file_path))

        # Read as bytes and extract
        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".docx", "test.docx")

        assert extracted is not None
        assert (
            "Test Document Title" in extracted or "test paragraph" in extracted
        )

    def test_xlsx_file_extraction(self, tmp_path):
        """Test extracting text from a real .xlsx file."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        from local_deep_research.document_loaders import extract_text_from_bytes

        # Create a real XLSX file
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["A2"] = "Test Item"
        ws["B2"] = 12345
        ws["A3"] = "Another Item"
        ws["B3"] = 67890

        file_path = tmp_path / "test.xlsx"
        wb.save(str(file_path))

        # Read as bytes and extract
        file_bytes = file_path.read_bytes()
        extracted = extract_text_from_bytes(file_bytes, ".xlsx", "test.xlsx")

        assert extracted is not None
        # Should contain some of the data
        assert (
            "Test Item" in extracted
            or "Name" in extracted
            or "12345" in extracted
        )

    def test_docx_from_path(self, tmp_path):
        """Test loading DOCX file using get_loader_for_path."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        from local_deep_research.document_loaders import get_loader_for_path

        # Create a real DOCX file
        doc = DocxDocument()
        doc.add_paragraph("Content loaded via path API")

        file_path = tmp_path / "path_test.docx"
        doc.save(str(file_path))

        loader = get_loader_for_path(file_path)

        assert loader is not None
        docs = loader.load()
        assert len(docs) > 0

    def test_xlsx_from_path(self, tmp_path):
        """Test loading XLSX file using get_loader_for_path."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        from local_deep_research.document_loaders import get_loader_for_path

        # Create a real XLSX file
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "Path Test Data"

        file_path = tmp_path / "path_test.xlsx"
        wb.save(str(file_path))

        loader = get_loader_for_path(file_path)

        assert loader is not None
        docs = loader.load()
        assert len(docs) > 0
